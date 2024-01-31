from typing import List, Any, Dict
import openpyxl
from docx import Document
from collections import defaultdict
from datetime import datetime
from docx.shared import Pt
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.enum.text import WD_UNDERLINE
n: datetime = datetime.now()
d: str = n.strftime("%d")
m: str = n.strftime("%m")
y: str = n.strftime("%y")
oati_count = caf_count = tot = co = 0
max_list: list[list[str]] = []
fin: list[Any] = []
accel: list[Any] = []
add_dict: dict[Any, Any] = {}
imprint: list[str] = []
head_cur: Any = []
single: list[Any] = []
if int(d[1])<=7 and int(d[0]) == 0:
    d2 = 30+int(d[1]) - 7
    m2 = m[0]+str(int(m[1]) - 1)
else:
    d2: int = int(d)-7
    m2: str = m
# Открываем файл Excel
workbook: Any = openpyxl.load_workbook('3.xlsx')
sheet: Any = workbook.active

# Создаем словарь для хранения количества
count_dict: dict[Any, int] = {}
count_dict2: dict[Any, int] = {}

# Итерируемся по строкам и считаем количество значений в колонке F, соответствующих значениям в колонке B
row: Any
for row in sheet.iter_rows(values_only=True):
    value_f: object = row[5]  # Индекс колонки F (нумерация начинается с 0)
    value_b: object = row[1]  # Индекс колонки B
    value_k: object = row[6] # Индекс колонки object element
    value_d: object = row[12] # Индекс колонки date
    value_ad: object = row[29] # Индекс колонки object address
    value_t: object = row[45] # Индекс статуса disturbance
    if value_t != 'Принятие мер исполнителем (Входящая)' and value_t != 'Завершено':
        continue
    if value_t == 'Принятие мер исполнителем (Входящая)' and value_b != None and value_b != 'ЦАФАП, ОАТИ' and value_b != 'Контрольный орган':
        accel.append(value_b)
        imprint.append(f" {value_f} - {value_ad}, {value_b}, Дата нарушения: {str(value_d).split(' ')[0]}; \n")
        imprint.append('\n')
    if value_f != 'Наименование нарушения' and value_f != None and value_t == 'Завершено':
        single.append(value_f)
        single.append(value_ad)
    if value_b == 'ЦАФАП':
        caf_count += 1
        if value_f in count_dict:
            count_dict[value_f] += 1
        else:
            count_dict[value_f] = 1
    if value_b == 'ОАТИ':
        oati_count += 1
        if value_f in count_dict2:
            count_dict2[value_f] += 1
        else:
            count_dict2[value_f] = 1
    if value_k == 'КП':
        uni_list = [f"{value_ad},{value_f}"]+[f"{str(value_d)}"]
        max_list.append(uni_list)

j: int
for j in range(len(max_list)-1):
    if max_list[j][0] == max_list[j+1][0]:
        fin.append(max_list[j])
        fin.append(max_list[j+1])


row_count: int = caf_count+oati_count

doc: Any = Document()
style: Any = doc.styles['Normal']
font: object = style.font
font.name = 'TimesNewRoman'
font.size = Pt(12)
e: Any = doc.add_paragraph('')
en: Any = e.add_run("Справка к повестке заседания окружного Штаба по вопросам жилищно-коммунального хозяйства и благоустройства \n \n")
en.font.bold = True
ent: Any = e.add_run("Вопрос №3 \n \n")
ent.font.bold = True
ent.font.underline = WD_UNDERLINE.SINGLE
run: object = e.add_run(f"О нарушениях, поступивших за истекшую неделю в ЦАФАП ГБУ «Жилищник района Теплый Стан» {d2}.{m2}.{y} г. - {d}.{m}.{y} г.")


head: Any = doc.add_paragraph(f"Всего поступило в работу {row_count} нарушений за период {d2}.{m2}.{y} г. - {d}.{m}.{y} г.")
head_1: Any  = head.add_run(f"\nОАТИ - {oati_count} нарушений \nЦАФАП - {caf_count} нарушений \n \n")
head_1.font.bold = True
head_2: Any = head.add_run("Из них по наименованиям нарушений:")
head_2.font.underline = WD_UNDERLINE.SINGLE
dd: defaultdict[Any, list] = defaultdict(list)

key: object
for key in set(list(count_dict.keys())+list(count_dict2.keys())):
    if key in count_dict:
        dd[key].append(str(count_dict[key]))
        dd[key].append(" (ЦАФАП) ")
    if key in count_dict2:
        dd[key].append(str(count_dict2[key]))
        dd[key].append(" (ОАТИ)")

count: str
for value_f, count in dd.items():
    tot = 0
    for j in range(len(count)):
        if count[j].isdigit():
            tot += int(count[j])
    count = ''.join(count)
    line: Any = doc.add_paragraph('')
    line_c: Any  = line.add_run(f"{value_f} - {count}, {tot} нарушений \n")
    line_c.font.bold = True
    line_1 = line.add_run("Нарушения устранены (находится на проверке)") 
    for i in range(len(single)-1):
        if tot == 1 and value_f == single[i]:
            doc.add_paragraph(single[i+1])
    for i in fin:
        if ''.join(i[0]).split(',')[1] == value_f:
            nai: Any = 'Наибольшее количество нарушений по адресам:\n'
            if co>=1:
                nai = None
            co +=1
            doc.add_paragraph(nai)
            add: str = ''.join(i[0]).split(',')[0]+' - '+''.join(i[1]).split(' ')[0]
            doc.add_paragraph(add)
        else:
            co = 0

if len(imprint)>0:
    doc.add_paragraph('\n')
    head_cur_0 = doc.add_paragraph('')
    head_cur = head_cur_0.add_run(f"В том числе всего незакрытых нарушений по состоянию на {d}.{m}.{y} г. – {len(accel)} нарушений, из них {accel.count('ОАТИ')} ОАТИ, {accel.count('ЦАФАП')} ЦАФАП\n")
    head_cur.font.bold  = True
    doc.add_paragraph(imprint)
doc.save('output.docx')

print("Документ успешно создан!")
